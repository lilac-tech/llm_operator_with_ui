import PyPDF2
import docx
"""
pdf_file = PyPDF2.PdfReader('ML code test (1).pdf')
num_pages = len(pdf_file.pages)
for page_num in range(num_pages):
    page = pdf_file.pages[page_num]
    text = page.extract_text()
    print(text)

document = docx.Document('ashwin_article.docx')
for para in document.paragraphs:
    print(para.text)
"""
class Reader:
    def __init__(self, file, type) -> None:
        self.reader = None
        self.new_file(file, type)

    def new_file(self, file, type):
        print(type)
        if type == r"application/pdf":
            self.reader = PdfReader(file)
            return
        if type == r"application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            self.reader = DocxReader(file)
            return
        self.reader = None
        print("No detected file type")
    
    def get_text(self):
        return self.reader.get_text()

class PdfReader:
    def __init__(self, file) -> None:
        self.pdf_file = PyPDF2.PdfReader(file)
        self.text = ""
        num_pages = len(self.pdf_file.pages)
        for page_num in range(num_pages):
            page = self.pdf_file.pages[page_num]
            self.text += page.extract_text()
    
    def get_text(self):
        return self.text

class DocxReader:
    def __init__(self, file):
        self.document = docx.Document(file)
        self.text = ""
        # Print the text of the document
        for para in self.document.paragraphs:
            self.text += para.text

    def get_text(self):
        return self.text